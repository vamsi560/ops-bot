import pandas as pd
from PyPDF2 import PdfReader
import os
import json
from typing import Dict, List, Any

# Utility functions to extract data from files

def extract_excel(file_path):
    """Extract data from Excel files with multiple sheets support"""
    try:
        excel_data = {}
        xl_file = pd.ExcelFile(file_path)
        
        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            excel_data[sheet_name] = {
                "data": df.to_dict(orient='records'),
                "columns": df.columns.tolist(),
                "shape": df.shape,
                "summary": {
                    "total_rows": len(df),
                    "total_columns": len(df.columns),
                    "column_types": df.dtypes.to_dict()
                }
            }
        
        return {
            "type": "excel",
            "sheets": excel_data,
            "total_sheets": len(excel_data)
        }
    except Exception as e:
        return {"error": f"Failed to extract Excel data: {str(e)}"}

def extract_csv(file_path):
    """Extract data from CSV files"""
    try:
        df = pd.read_csv(file_path)
        return {
            "type": "csv",
            "data": df.to_dict(orient='records'),
            "columns": df.columns.tolist(),
            "shape": df.shape,
            "summary": {
                "total_rows": len(df),
                "total_columns": len(df.columns),
                "column_types": df.dtypes.to_dict()
            }
        }
    except Exception as e:
        return {"error": f"Failed to extract CSV data: {str(e)}"}

def extract_pdf(file_path):
    """Extract text from PDF files"""
    try:
        reader = PdfReader(file_path)
        text_content = []
        page_data = []
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text_content.append(text)
            page_data.append({
                "page_number": page_num + 1,
                "text": text,
                "char_count": len(text)
            })
        
        return {
            "type": "pdf",
            "text": "\n".join(text_content),
            "pages": page_data,
            "total_pages": len(reader.pages),
            "total_characters": len("\n".join(text_content))
        }
    except Exception as e:
        return {"error": f"Failed to extract PDF data: {str(e)}"}

def extract_word(file_path):
    """Extract text from Word files"""
    try:
        # Try to use python-docx if available
        try:
            from docx import Document
            doc = Document(file_path)
            text_content = []
            paragraph_data = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_content.append(text)
                    paragraph_data.append({
                        "text": text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables if any
            table_data = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                table_data.append(table_rows)
            
            return {
                "type": "word",
                "text": "\n".join(text_content),
                "paragraphs": paragraph_data,
                "tables": table_data,
                "total_paragraphs": len(paragraph_data),
                "total_tables": len(table_data),
                "total_characters": len("\n".join(text_content))
            }
        except ImportError:
            return {
                "type": "word",
                "message": "Word extraction requires python-docx library",
                "suggestion": "Install python-docx for full Word support",
                "file_path": file_path
            }
    except Exception as e:
        return {"error": f"Failed to extract Word data: {str(e)}"}

def extract_powerpoint(file_path):
    """Extract text from PowerPoint files"""
    try:
        # Note: This is a basic implementation. For production, consider using python-pptx
        # For now, we'll return a placeholder structure
        return {
            "type": "powerpoint",
            "message": "PowerPoint extraction requires python-pptx library",
            "suggestion": "Install python-pptx for full PowerPoint support",
            "file_path": file_path
        }
    except Exception as e:
        return {"error": f"Failed to extract PowerPoint data: {str(e)}"}

def extract_file(file_path):
    """Main function to extract data from various file types"""
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    ext = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)
    
    extraction_result = {
        "file_name": file_name,
        "file_path": file_path,
        "file_extension": ext,
        "file_size": os.path.getsize(file_path)
    }
    
    try:
        if ext in [".xlsx", ".xls"]:
            data = extract_excel(file_path)
        elif ext == ".csv":
            data = extract_csv(file_path)
        elif ext == ".pdf":
            data = extract_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            data = extract_word(file_path)
        elif ext in [".pptx", ".ppt"]:
            data = extract_powerpoint(file_path)
        else:
            data = {"error": f"Unsupported file type: {ext}"}
        
        extraction_result.update(data)
        return extraction_result
        
    except Exception as e:
        extraction_result["error"] = f"Extraction failed: {str(e)}"
        return extraction_result

def get_file_summary(file_path):
    """Get a quick summary of file contents without full extraction"""
    try:
        ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        if ext in [".xlsx", ".xls"]:
            xl_file = pd.ExcelFile(file_path)
            return {
                "file_name": file_name,
                "type": "excel",
                "sheets": xl_file.sheet_names,
                "sheet_count": len(xl_file.sheet_names)
            }
        elif ext == ".csv":
            df = pd.read_csv(file_path, nrows=5)  # Read only first 5 rows for summary
            return {
                "file_name": file_name,
                "type": "csv",
                "columns": df.columns.tolist(),
                "column_count": len(df.columns),
                "sample_rows": df.to_dict(orient='records')
            }
        elif ext == ".pdf":
            reader = PdfReader(file_path)
            return {
                "file_name": file_name,
                "type": "pdf",
                "page_count": len(reader.pages)
            }
        else:
            return {
                "file_name": file_name,
                "type": "unknown",
                "extension": ext
            }
    except Exception as e:
        return {
            "file_name": os.path.basename(file_path),
            "error": str(e)
        }
